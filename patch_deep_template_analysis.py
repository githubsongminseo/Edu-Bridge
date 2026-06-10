#!/usr/bin/env python3
"""
patch_deep_template_analysis.py
================================
양식 업로드 시 Gemini로 심층 분석 + 위치 기반 채우기

핵심 개선:
1. 업로드 시: 양식의 모든 셀 위치와 내용 분석 → Gemini가 각 셀의 의미 파악
2. 다운로드 시: 분석 결과를 사용해 정확한 셀 위치에 내용 채우기
3. 복잡한 병합 셀과 다단 구조도 정확히 처리
"""
from pathlib import Path

BACKEND = Path(".")
SERVICES = BACKEND / "services"
MAIN_PATH = BACKEND / "main.py"
DB_PATH = BACKEND / "database.py"


# ============================================================
# 1) services/template_analyzer.py 신규 생성
# ============================================================
ANALYZER_CODE = '''"""
services/template_analyzer.py
==============================
양식을 셀 단위로 분석하고 Gemini로 각 셀의 의미를 파악
"""
import json
from typing import List, Dict
from docx import Document


def extract_template_cells(file_path: str) -> List[Dict]:
    """
    양식의 모든 셀을 위치 정보와 함께 추출 (병합 셀 처리)

    Returns:
        [
            {"table_idx": 0, "row": 0, "col": 0, "text": "단원(차시)", "is_empty": False},
            {"table_idx": 0, "row": 0, "col": 1, "text": "", "is_empty": True},
            ...
        ]
    """
    doc = Document(file_path)
    cells_info = []

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen_in_row = set()  # 행 내 병합 셀 추적
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                # 병합 셀: 같은 셀 객체가 반복되면 첫 번째만 사용
                cell_key = id(cell._tc)
                if cell_key in seen_in_row:
                    continue
                seen_in_row.add(cell_key)

                cells_info.append({
                    "table_idx": ti,
                    "row": ri,
                    "col": ci,
                    "text": text[:200],  # 너무 긴 텍스트는 자르기
                    "is_empty": not text or _is_placeholder(text),
                })

    return cells_info


def _is_placeholder(text: str) -> bool:
    """예시/플레이스홀더 텍스트인지 판단"""
    if not text:
        return True
    text = text.strip()
    if len(text) < 2:
        return True
    placeholders = ["...", "___", "OOO", "○○○", "(예", "예)", "예시"]
    return any(p in text for p in placeholders)


def analyze_template_deeply(file_path: str) -> Dict:
    """
    Gemini를 사용해 양식의 각 셀을 분석.

    Returns:
        {
            "cells": [
                {
                    "table_idx": 0, "row": 0, "col": 0,
                    "role": "label",  # "label" or "data"
                    "category": "lesson_meta",  # personal_info, lesson_meta, lesson_content, evaluation
                    "label_text": "단원(차시)",
                },
                {
                    "table_idx": 0, "row": 0, "col": 1,
                    "role": "data",
                    "for_label": "단원(차시)",
                    "category": "lesson_meta",
                    "expected_content": "단원의 이름과 차시 (예: 2단원 모양 알기 (3/8))",
                    "approx_length": "short",  # short / medium / long
                },
                ...
            ],
            "summary": "유아 활동 지도안 양식 - 표 6개로 구성. 개인정보, 수업 메타, 수업 내용, 평가 영역 포함",
        }
    """
    cells_info = extract_template_cells(file_path)
    if not cells_info:
        return {"cells": [], "summary": "표가 없는 양식 (단락 기반)"}

    # 셀 정보를 Gemini에게 보낼 텍스트로 정리
    cells_text = []
    for c in cells_info:
        marker = "[비어있음]" if c["is_empty"] else c["text"]
        cells_text.append(f"표{c['table_idx']} 행{c['row']} 열{c['col']}: {marker}")

    cells_str = "\\n".join(cells_text)

    prompt = f"""아래는 한국 유아교육 지도안 양식의 모든 셀 정보입니다.
각 셀을 분석해서 역할과 의미를 파악해주세요.

## 셀 목록
{cells_str}

## 작업 지침
각 셀에 대해 다음을 판단하세요:

1. **role**: "label" (헤더/라벨) 또는 "data" (값이 들어갈 셀)
2. **category** (label인 경우):
   - "personal_info": 교사명, 학생명, 학교명, 학과, 결재란 등 사람/조직 정보
   - "lesson_meta": 단원, 차시, 학습 주제, 학습 목표, 성취 기준, 교과 역량 등 수업 정보
   - "lesson_content": 학습 단계, 도입, 전개, 정리, 교사 활동, 학생 활동, 유의점 등 수업 내용
   - "evaluation": 평가 기준, 평가 방법, 성취 수준 등 평가 영역
   - "resource": 교수학습자료, 준비물 등
3. **for_label** (data인 경우): 어떤 라벨에 대한 데이터인지
4. **expected_content** (data인 경우): 어떤 내용이 들어가야 하는지 구체적으로 설명
5. **approx_length** (data인 경우): "short" (한 줄), "medium" (몇 문장), "long" (한 단락 이상)

## 출력 형식
오직 JSON만 출력. 다른 설명 없이:
{{
  "cells": [
    {{"table_idx": 0, "row": 0, "col": 0, "role": "label", "category": "lesson_meta", "label_text": "단원(차시)"}},
    {{"table_idx": 0, "row": 0, "col": 1, "role": "data", "for_label": "단원(차시)", "category": "lesson_meta", "expected_content": "단원 이름과 차시 표기 (예: 2단원 모양 알기 3/8차시)", "approx_length": "short"}},
    ...
  ],
  "summary": "이 양식에 대한 1-2줄 설명"
}}

모든 셀에 대해 분석하세요. 단, 같은 텍스트가 반복되는 병합 셀은 한 번만 분석.
빈 셀이라도 어떤 라벨의 데이터인지 추론해서 role을 "data"로 분류하고 for_label과 expected_content를 채워주세요.
"""

    try:
        from services.keyword_extractor import _get_client
        from google.genai import types as genai_types

        client = _get_client()
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                temperature=0.2,
                max_output_tokens=8192,
                response_mime_type="application/json",
                thinking_config=genai_types.ThinkingConfig(thinking_budget=0),
            ),
        )

        text = response.text.strip()
        if text.startswith("```"):
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:].strip()
            text = text.split("```")[0].strip()

        result = json.loads(text)
        if "cells" not in result:
            result["cells"] = []
        return result
    except Exception as e:
        print(f"⚠️ 양식 심층 분석 실패: {e}")
        # 실패해도 빈 결과 반환 (기본 채우기 폴백)
        return {"cells": [], "summary": f"분석 실패: {e}", "error": str(e)}


def fill_cells_with_lesson(
    template_analysis: Dict,
    lesson_markdown: str,
    lesson_meta: Dict = None,
    personal_info: Dict = None,
) -> Dict:
    """
    심층 분석된 양식 정보 + 지도안 마크다운 → 각 데이터 셀에 들어갈 내용 매핑

    Returns:
        {
            "(table_idx, row, col)": "채울 내용",
            ...
        }
    """
    cells = template_analysis.get("cells", [])
    data_cells = [c for c in cells if c.get("role") == "data"]

    if not data_cells:
        return {}

    # 개인정보 셀과 수업 내용 셀 분리
    # personal_info에 매칭되는 셀은 사용자 입력값 사용
    direct_fills = {}  # (ti,ri,ci) -> content
    cells_for_gemini = []

    for c in data_cells:
        position = (c["table_idx"], c["row"], c["col"])
        category = c.get("category", "")
        for_label = c.get("for_label", "")

        # 개인정보 셀: 사용자 입력에서 직접 매칭
        if category == "personal_info" and personal_info:
            matched = None
            for label, value in personal_info.items():
                if label in for_label or for_label in label:
                    matched = value
                    break
            if matched:
                direct_fills[position] = matched
                continue

        # 나머지는 Gemini에게 매핑 요청
        cells_for_gemini.append(c)

    # Gemini로 수업 내용 셀 채우기
    if cells_for_gemini:
        gemini_fills = _gemini_fill_data_cells(
            cells_for_gemini, lesson_markdown, lesson_meta or {}
        )
        for position, content in gemini_fills.items():
            if position not in direct_fills:
                direct_fills[position] = content

    return direct_fills


def _gemini_fill_data_cells(
    data_cells: List[Dict],
    lesson_markdown: str,
    lesson_meta: Dict,
) -> Dict:
    """데이터 셀들에 대해 Gemini로 내용 생성"""
    if not data_cells:
        return {}

    # 셀 목록을 ID와 함께 정리
    cells_for_prompt = []
    for i, c in enumerate(data_cells):
        cells_for_prompt.append({
            "id": i,
            "label": c.get("for_label", "(unknown)"),
            "category": c.get("category", ""),
            "expected": c.get("expected_content", ""),
            "length": c.get("approx_length", "medium"),
        })

    meta_str = ""
    if lesson_meta:
        parts = []
        if "age" in lesson_meta:
            parts.append(f"대상 연령: 만 {lesson_meta['age']}세")
        if "duration" in lesson_meta:
            parts.append(f"수업 시간: {lesson_meta['duration']}분")
        if "search_query" in lesson_meta:
            parts.append(f"수업 주제: {lesson_meta['search_query']}")
        if parts:
            meta_str = "\\n## 수업 메타 정보\\n" + "\\n".join(parts)

    prompt = f"""아래는 [생성된 지도안]과 [채워야 할 양식 셀 목록]입니다.
각 셀에 들어갈 내용을 지도안 내용을 바탕으로 작성해주세요.

## 생성된 지도안 (마크다운)
{lesson_markdown}
{meta_str}

## 채워야 할 양식 셀 목록
{json.dumps(cells_for_prompt, ensure_ascii=False, indent=2)}

## 작업 지침
1. 각 셀의 label과 expected_content를 보고 지도안에서 적절한 내용을 추출/요약하세요.
2. length가 "short"면 한 줄 이내, "medium"이면 1-3 문장, "long"이면 한 단락으로.
3. 마크다운 기호(#, *, -)는 제거하고 깔끔한 문장으로.
4. 지도안에 직접적인 정보가 없는 경우, 수업 메타 정보와 expected_content 설명에 기반해 합리적으로 작성.
5. 평가/성취기준 등은 일반적인 유아교육 기준에 맞춰 적절히 생성.
6. 학습 단계(도입/전개/정리)별 셀이면 지도안의 해당 단계 내용을 정확히 매핑.
7. 모든 셀에 내용을 채워주세요. 빈 값 없이.

## 출력 형식
오직 JSON만 출력:
{{
  "fills": [
    {{"id": 0, "content": "..."}},
    {{"id": 1, "content": "..."}},
    ...
  ]
}}
"""

    try:
        from services.keyword_extractor import _get_client
        from google.genai import types as genai_types

        client = _get_client()
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                temperature=0.4,
                max_output_tokens=8192,
                response_mime_type="application/json",
                thinking_config=genai_types.ThinkingConfig(thinking_budget=0),
            ),
        )

        text = response.text.strip()
        if text.startswith("```"):
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:].strip()
            text = text.split("```")[0].strip()

        result = json.loads(text)
        fills = result.get("fills", [])

        # id → (ti, ri, ci) 매핑
        position_fills = {}
        for fill in fills:
            idx = fill.get("id")
            content = fill.get("content", "")
            if idx is not None and 0 <= idx < len(data_cells) and content:
                c = data_cells[idx]
                position = (c["table_idx"], c["row"], c["col"])
                position_fills[position] = content

        return position_fills
    except Exception as e:
        print(f"⚠️ Gemini 셀 채우기 실패: {e}")
        return {}
'''

(SERVICES / "template_analyzer.py").write_text(ANALYZER_CODE, encoding="utf-8")
print("✅ services/template_analyzer.py 생성")


# ============================================================
# 2) docx_writer.py에 위치 기반 채우기 함수 추가
# ============================================================
docx_writer = (SERVICES / "docx_writer.py").read_text(encoding="utf-8")

new_func = '''


def fill_template_by_positions(template_path: str, output_path: str, position_fills: dict):
    """
    셀 위치 기반으로 양식 채우기.

    Args:
        position_fills: {(table_idx, row, col): "content", ...}
    """
    doc = Document(template_path)

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                key = (table_idx, row_idx, col_idx)
                if key in position_fills:
                    content = position_fills[key]
                    if content:
                        _replace_cell_text(cell, str(content))

    doc.save(output_path)


def markdown_to_docx_with_analysis(
    markdown: str,
    output_path: str,
    title: str,
    template_path: str,
    template_analysis: dict,
    lesson_meta: dict = None,
    personal_info: dict = None,
):
    """
    심층 분석된 양식 정보로 docx 생성
    """
    from services.template_analyzer import fill_cells_with_lesson

    position_fills = fill_cells_with_lesson(
        template_analysis=template_analysis,
        lesson_markdown=markdown,
        lesson_meta=lesson_meta,
        personal_info=personal_info,
    )

    print(f"📊 채울 셀 위치: {len(position_fills)}개")

    fill_template_by_positions(template_path, output_path, position_fills)
'''

if "fill_template_by_positions" not in docx_writer:
    docx_writer += new_func
    (SERVICES / "docx_writer.py").write_text(docx_writer, encoding="utf-8")
    print("✅ docx_writer.py: fill_template_by_positions + markdown_to_docx_with_analysis 추가")


# ============================================================
# 3) database.py에 analysis_json 컬럼 추가
# ============================================================
db_code = DB_PATH.read_text(encoding="utf-8")

if "analysis_json" not in db_code:
    old = '    structure_type = Column(String(50), nullable=True)  # "table" or "paragraph"'
    new = '''    structure_type = Column(String(50), nullable=True)  # "table" or "paragraph"
    analysis_json = Column(Text, nullable=True)  # Gemini 심층 분석 결과 (셀 위치별 의미)'''
    if old in db_code:
        db_code = db_code.replace(old, new)
        DB_PATH.write_text(db_code, encoding="utf-8")
        print("✅ database.py: analysis_json 컬럼 추가")
        print("⚠️  rm edubridge.db 필요 (스키마 변경)")


# ============================================================
# 4) main.py 업로드 API에 심층 분석 추가
# ============================================================
main_code = MAIN_PATH.read_text(encoding="utf-8")

# 업로드 시점에 심층 분석 호출
old_upload = '''        # DB 저장
        tpl = UserTemplate(
            user_id=current_user.id,
            template_name=template_name,
            original_filename=file.filename,
            file_path=str(saved_path),
            sections_json=json.dumps(parsed["sections"], ensure_ascii=False),
            structure_type=parsed["structure_type"],
        )'''

new_upload = '''        # 심층 분석 (Gemini로 셀별 의미 파악, .docx만)
        analysis_data = None
        if not is_hwp:
            try:
                from services.template_analyzer import analyze_template_deeply
                analysis_data = analyze_template_deeply(str(saved_path))
                print(f"🔍 양식 심층 분석 완료: {len(analysis_data.get('cells', []))}개 셀")
            except Exception as e:
                print(f"⚠️ 심층 분석 실패 (계속 진행): {e}")

        # DB 저장
        tpl = UserTemplate(
            user_id=current_user.id,
            template_name=template_name,
            original_filename=file.filename,
            file_path=str(saved_path),
            sections_json=json.dumps(parsed["sections"], ensure_ascii=False),
            structure_type=parsed["structure_type"],
            analysis_json=json.dumps(analysis_data, ensure_ascii=False) if analysis_data else None,
        )'''

if old_upload in main_code:
    main_code = main_code.replace(old_upload, new_upload)
    print("✅ main.py: 업로드 시 심층 분석 호출 추가")
else:
    print("⚠️  main.py: 업로드 패턴 못 찾음")


# export-docx에서 분석 결과 활용
old_export_logic = '''                if tpl_sections:
                    try:
                        from services.template_filler import smart_fill_with_gemini
                        lesson_meta = {
                            "age": body.get("age"),
                            "duration": body.get("duration"),
                            "search_query": body.get("search_query"),
                        }
                        lesson_meta = {k: v for k, v in lesson_meta.items() if v is not None}
                        sections_data = smart_fill_with_gemini(
                            lesson_markdown=markdown,
                            template_sections=tpl_sections,
                            lesson_meta=lesson_meta,
                        )
                        print(f"✅ Gemini 매핑 완료: {len(sections_data)}개 섹션")
                    except Exception as e:
                        print(f"⚠️ Gemini 매핑 실패, 기본 채우기로 폴백: {e}")
                        sections_data = body.get("sections_data")
                    
                    # 사용자가 입력한 개인 정보로 덮어쓰기 (Gemini가 추측한 값보다 우선)
                    personal_info = body.get("personal_info") or {}
                    if personal_info and isinstance(sections_data, dict):
                        for label, value in personal_info.items():
                            if value and value.strip():
                                sections_data[label] = value.strip()
                        print(f"✅ 개인정보 {len(personal_info)}개 적용")'''

new_export_logic = '''                # 양식의 심층 분석 결과가 있으면 위치 기반 채우기
                template_analysis = None
                if tpl.analysis_json:
                    try:
                        template_analysis = json.loads(tpl.analysis_json)
                    except Exception as e:
                        print(f"⚠️ 분석 JSON 파싱 실패: {e}")

                lesson_meta = {
                    "age": body.get("age"),
                    "duration": body.get("duration"),
                    "search_query": body.get("search_query"),
                }
                lesson_meta = {k: v for k, v in lesson_meta.items() if v is not None}
                personal_info = body.get("personal_info") or {}

                if template_analysis and template_analysis.get("cells"):
                    # 위치 기반 채우기 사용
                    from services.docx_writer import markdown_to_docx_with_analysis
                    try:
                        markdown_to_docx_with_analysis(
                            markdown=markdown,
                            output_path=str(output_path),
                            title=title,
                            template_path=template_path,
                            template_analysis=template_analysis,
                            lesson_meta=lesson_meta,
                            personal_info=personal_info,
                        )
                        filename_for_download = f"{title}.docx".replace("/", "_").replace(" ", "_")
                        return FileResponse(
                            path=str(output_path),
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            filename=filename_for_download,
                        )
                    except Exception as e:
                        print(f"⚠️ 위치 기반 채우기 실패, 기존 방식으로 폴백: {e}")

                # 폴백: 기존 이름 기반 매핑
                if tpl_sections:
                    try:
                        from services.template_filler import smart_fill_with_gemini
                        sections_data = smart_fill_with_gemini(
                            lesson_markdown=markdown,
                            template_sections=tpl_sections,
                            lesson_meta=lesson_meta,
                        )
                        print(f"✅ Gemini 매핑 완료: {len(sections_data)}개 섹션")
                    except Exception as e:
                        print(f"⚠️ Gemini 매핑 실패, 기본 채우기로 폴백: {e}")
                        sections_data = body.get("sections_data")

                    if personal_info and isinstance(sections_data, dict):
                        for label, value in personal_info.items():
                            if value and value.strip():
                                sections_data[label] = value.strip()
                        print(f"✅ 개인정보 {len(personal_info)}개 적용")'''

if old_export_logic in main_code:
    main_code = main_code.replace(old_export_logic, new_export_logic)
    print("✅ main.py: export-docx에 위치 기반 채우기 추가")
else:
    print("⚠️  main.py: export-docx 패턴 못 찾음")

MAIN_PATH.write_text(main_code, encoding="utf-8")

print("\n" + "=" * 50)
print("🎉 양식 심층 분석 패치 완료!")
print("=" * 50)
print("\n다음 단계:")
print("  1. rm edubridge.db  (DB 스키마 변경)")
print("  2. 서버 재시작 (이미 자동 재시작됨)")
print("  3. 양식 다시 업로드 (분석 데이터 저장됨)")
print("  4. 지도안 생성 → 양식 적용 .docx 다운로드")
print("\n변경된 흐름:")
print("  업로드: 양식 → Gemini 셀별 의미 분석 → DB에 위치+의미 저장")
print("  다운로드: Gemini가 지도안 내용을 정확한 셀 위치에 매핑 → 채움")
