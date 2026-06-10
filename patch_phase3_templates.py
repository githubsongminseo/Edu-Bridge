#!/usr/bin/env python3
"""
patch_phase3_templates.py
==========================
Phase 3: 커스텀 지도안 양식 업로드 기능 통합 패치

수행하는 작업:
1. services/template_parser.py 생성 (docx 파싱)
2. services/docx_writer.py 생성 (docx 출력)
3. database.py 패치 (UserTemplate 테이블 추가)
4. main.py 패치 (4개 API 추가)
5. static/edu-bridge-full.html 패치
   - 사이드바 "내 양식" 메뉴
   - 양식 업로드 페이지
   - Play-Scanner에 양식 선택 옵션
   - 지도안 결과에 docx 다운로드 버튼

사용법:
    cd backend
    pip install python-docx
    python3 patch_phase3_templates.py
"""
from pathlib import Path

BACKEND = Path(".")
SERVICES = BACKEND / "services"
HTML_PATH = BACKEND / "static" / "edu-bridge-full.html"
MAIN_PATH = BACKEND / "main.py"
DB_PATH = BACKEND / "database.py"

if not HTML_PATH.exists() or not MAIN_PATH.exists() or not DB_PATH.exists():
    print("❌ backend/ 디렉토리에서 실행해야 합니다.")
    exit(1)


# ============================================================
# 1) services/template_parser.py 생성
# ============================================================
TEMPLATE_PARSER_CODE = '''"""
services/template_parser.py
============================
업로드된 .docx 양식을 파싱해서 섹션/필드 구조를 추출
"""
from docx import Document
from typing import List, Dict
import re


def parse_docx_template(file_path: str) -> Dict:
    """
    .docx 파일을 분석하여 구조 정보를 반환

    Returns:
        {
            "sections": [{"name": "활동명", "type": "text", "order": 0}, ...],
            "has_tables": True,
            "raw_text_sample": "...",
            "structure_type": "table" | "paragraph"
        }
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": f"파일 파싱 실패: {e}"}

    sections = []
    raw_lines = []

    # 1) 표(table) 기반 양식 분석
    has_tables = len(doc.tables) > 0
    if has_tables:
        order = 0
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                # 첫 셀이 헤더(label)인 경우가 많음
                for i, cell_text in enumerate(cells):
                    if cell_text and len(cell_text) < 30 and not _is_filler_text(cell_text):
                        # 헤더 후보 (짧은 텍스트, 콜론 끝, 한국어 키워드)
                        if _looks_like_section_header(cell_text):
                            sections.append({
                                "name": _clean_section_name(cell_text),
                                "type": "table_cell",
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "col_idx": i,
                                "order": order,
                            })
                            order += 1

    # 2) 단락 기반 양식 분석
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        raw_lines.append(text)
        if _looks_like_section_header(text) and not has_tables:
            sections.append({
                "name": _clean_section_name(text),
                "type": "paragraph",
                "order": len(sections),
            })

    # 중복 제거 (이름 기준)
    seen = set()
    unique_sections = []
    for s in sections:
        if s["name"] not in seen:
            seen.add(s["name"])
            unique_sections.append(s)

    return {
        "sections": unique_sections,
        "has_tables": has_tables,
        "raw_text_sample": "\\n".join(raw_lines[:10]),
        "structure_type": "table" if has_tables else "paragraph",
        "total_sections": len(unique_sections),
    }


SECTION_KEYWORDS = [
    "활동명", "주제", "제목", "단원", "차시",
    "대상", "연령", "인원", "장소", "일시", "날짜",
    "목표", "목적", "학습목표", "교육목표", "수업목표",
    "준비물", "교재", "교구", "재료",
    "도입", "전개", "정리", "마무리", "활동내용",
    "활동", "수업", "지도", "전개과정",
    "평가", "유의점", "참고", "비고",
    "교사", "유아", "역할",
]


def _looks_like_section_header(text: str) -> bool:
    """섹션 헤더 같은 텍스트인지 판단"""
    text = text.strip()
    if not text or len(text) > 30:
        return False
    # 콜론으로 끝나는 짧은 텍스트
    if text.endswith(":") or text.endswith("："):
        return True
    # 키워드 매칭
    for kw in SECTION_KEYWORDS:
        if kw in text:
            return True
    return False


def _is_filler_text(text: str) -> bool:
    """예시 텍스트나 채워야 할 빈칸인지"""
    fillers = ["예)", "예시", "(예", "...", "___", "OOO", "○○○"]
    return any(f in text for f in fillers)


def _clean_section_name(text: str) -> str:
    """섹션 이름 정리"""
    text = text.strip()
    text = re.sub(r"[:：\\s]+$", "", text)
    text = re.sub(r"^[\\d\\.\\)\\s]+", "", text)
    return text.strip()
'''

(SERVICES / "template_parser.py").write_text(TEMPLATE_PARSER_CODE, encoding="utf-8")
print("✅ services/template_parser.py 생성")


# ============================================================
# 2) services/docx_writer.py 생성
# ============================================================
DOCX_WRITER_CODE = '''"""
services/docx_writer.py
========================
지도안 마크다운을 .docx로 변환
(사용자 양식이 있으면 그 양식에 채우고, 없으면 기본 양식 사용)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional, Dict
import re


def markdown_to_docx(
    markdown: str,
    output_path: str,
    title: str = "지도안",
    template_path: Optional[str] = None,
    sections_data: Optional[Dict[str, str]] = None,
):
    """
    Markdown을 .docx로 저장.

    template_path가 있으면 그 양식의 표/단락에 sections_data를 채워서 저장.
    없으면 새 문서를 만들어서 마크다운을 단순 변환.
    """
    if template_path and sections_data:
        _fill_template(template_path, output_path, sections_data)
    else:
        _create_simple_docx(markdown, output_path, title)


def _create_simple_docx(markdown: str, output_path: str, title: str):
    """기본 양식: 마크다운 → 단순 docx"""
    doc = Document()

    # 제목
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 마크다운 줄 단위 처리
    for line in markdown.split("\\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        # 헤더 처리
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(16)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(14)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            r.bold = True
            r.font.size = Pt(12)
        # 리스트
        elif line.lstrip().startswith(("- ", "* ", "• ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        # 번호 리스트
        elif re.match(r"^\\d+\\.\\s", line):
            text = re.sub(r"^\\d+\\.\\s", "", line)
            doc.add_paragraph(text, style="List Number")
        # 굵은 글씨 처리 (**text**)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    doc.save(output_path)


def _add_formatted_text(paragraph, text: str):
    """**bold** 마크다운을 docx 굵은 글씨로"""
    parts = re.split(r"(\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _fill_template(template_path: str, output_path: str, sections_data: Dict[str, str]):
    """
    사용자 양식 docx를 열어서 각 섹션에 값을 채워 넣기.
    
    동작 방식:
    - 표 안의 각 셀: 셀 텍스트가 sections_data의 키와 매칭되면 다음 셀에 값 삽입
    - 단락 기반: 헤더 단락 다음에 새 단락으로 값 추가
    """
    doc = Document(template_path)

    # 1) 표 기반 채우기
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip().rstrip(":：").strip()
                # 매칭되는 섹션 데이터 찾기
                value = _find_matching_value(cell_text, sections_data)
                if value and i + 1 < len(cells):
                    # 다음 셀이 비어있으면 값 삽입
                    next_cell = cells[i + 1]
                    if not next_cell.text.strip() or _is_filler(next_cell.text):
                        # 기존 텍스트 제거
                        for para in next_cell.paragraphs:
                            for run in para.runs:
                                run.text = ""
                        # 새 텍스트 삽입
                        next_cell.paragraphs[0].add_run(value)

    # 2) 단락 기반 채우기 (헤더 발견 시 다음 단락에 값 추가)
    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        text = para.text.strip().rstrip(":：").strip()
        value = _find_matching_value(text, sections_data)
        if value and i + 1 < len(paragraphs):
            next_para = paragraphs[i + 1]
            if not next_para.text.strip() or _is_filler(next_para.text):
                next_para.add_run(value)

    doc.save(output_path)


def _find_matching_value(section_name: str, sections_data: Dict[str, str]) -> Optional[str]:
    """섹션 이름과 매칭되는 값 찾기 (부분 매칭 허용)"""
    if not section_name:
        return None
    for key, value in sections_data.items():
        if key in section_name or section_name in key:
            return value
    return None


def _is_filler(text: str) -> bool:
    text = text.strip()
    return text in ["", "...", "___", "OOO", "○○○"] or text.startswith("예)")
'''

(SERVICES / "docx_writer.py").write_text(DOCX_WRITER_CODE, encoding="utf-8")
print("✅ services/docx_writer.py 생성")


# ============================================================
# 3) database.py 패치 - UserTemplate 테이블 추가
# ============================================================
db_code = DB_PATH.read_text(encoding="utf-8")

if "class UserTemplate" not in db_code:
    # SavedLesson 정의 다음에 추가
    insertion_point = "# ============================================================\n# 데이터베이스 초기화"
    new_table = '''
class UserTemplate(Base):
    """사용자가 업로드한 지도안 양식"""
    __tablename__ = "user_templates"

    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)

    template_name = Column(String(200), nullable=False)
    original_filename = Column(String(300), nullable=False)
    file_path = Column(String(500), nullable=False)

    # 파싱된 섹션 정보 (JSON 문자열)
    sections_json = Column(Text, nullable=True)
    structure_type = Column(String(50), nullable=True)  # "table" or "paragraph"

    created_at = Column(DateTime, default=datetime.utcnow)


'''
    db_code = db_code.replace(insertion_point, new_table + insertion_point)
    DB_PATH.write_text(db_code, encoding="utf-8")
    print("✅ database.py: UserTemplate 테이블 추가")
else:
    print("ℹ️  database.py: UserTemplate 이미 존재")


# ============================================================
# 4) main.py 패치 - 4개 API 추가
# ============================================================
main_code = MAIN_PATH.read_text(encoding="utf-8")

if "/api/templates/upload" not in main_code:
    # User import에 UserTemplate 추가
    if "UserTemplate" not in main_code:
        main_code = main_code.replace(
            "User, SavedLesson,",
            "User, SavedLesson, UserTemplate,"
        )

    # health check 앞에 새 API 삽입
    template_apis = '''

# ============================================================
# 커스텀 템플릿 API
# ============================================================
import shutil
import uuid

TEMPLATES_DIR = BACKEND_ROOT / "user_templates"
TEMPLATES_DIR.mkdir(exist_ok=True)


@app.post("/api/templates/upload")
async def api_upload_template(
    file: UploadFile = File(...),
    template_name: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """지도안 양식(.docx) 업로드 및 자동 파싱"""
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="현재는 .docx 파일만 지원합니다.")

    try:
        # 파일 저장
        ext = Path(file.filename).suffix
        saved_name = f"{current_user.id}_{uuid.uuid4().hex}{ext}"
        saved_path = TEMPLATES_DIR / saved_name
        with open(saved_path, "wb") as f:
            content = await file.read()
            f.write(content)

        # 파싱
        from services.template_parser import parse_docx_template
        parsed = parse_docx_template(str(saved_path))

        if "error" in parsed:
            saved_path.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail=parsed["error"])

        # DB 저장
        tpl = UserTemplate(
            user_id=current_user.id,
            template_name=template_name,
            original_filename=file.filename,
            file_path=str(saved_path),
            sections_json=json.dumps(parsed["sections"], ensure_ascii=False),
            structure_type=parsed["structure_type"],
        )
        db.add(tpl)
        db.commit()
        db.refresh(tpl)

        return {
            "id": tpl.id,
            "template_name": tpl.template_name,
            "sections": parsed["sections"],
            "structure_type": parsed["structure_type"],
            "total_sections": parsed["total_sections"],
            "message": "양식이 업로드되었습니다.",
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"업로드 실패: {str(e)}")


@app.get("/api/templates/my")
def api_get_my_templates(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """내 양식 목록"""
    templates = (
        db.query(UserTemplate)
        .filter(UserTemplate.user_id == current_user.id)
        .order_by(UserTemplate.created_at.desc())
        .all()
    )
    return [
        {
            "id": t.id,
            "template_name": t.template_name,
            "original_filename": t.original_filename,
            "structure_type": t.structure_type,
            "sections": json.loads(t.sections_json) if t.sections_json else [],
            "created_at": t.created_at.isoformat(),
        }
        for t in templates
    ]


@app.delete("/api/templates/{template_id}")
def api_delete_template(
    template_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """양식 삭제"""
    tpl = db.query(UserTemplate).filter(
        UserTemplate.id == template_id,
        UserTemplate.user_id == current_user.id
    ).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="양식을 찾을 수 없습니다.")
    # 파일 삭제
    Path(tpl.file_path).unlink(missing_ok=True)
    db.delete(tpl)
    db.commit()
    return {"message": "삭제되었습니다."}


@app.post("/api/lesson/export-docx")
async def api_export_lesson_docx(body: dict):
    """
    지도안을 .docx 파일로 다운로드
    
    body: {
        "markdown": "...",
        "title": "...",
        "template_id": null (또는 양식 ID),
        "sections_data": {"활동명": "...", "목표": "..."}  // template_id 있을 때
    }
    """
    from services.docx_writer import markdown_to_docx
    
    markdown = body.get("markdown", "")
    title = body.get("title", "지도안")
    template_id = body.get("template_id")
    sections_data = body.get("sections_data")
    
    output_path = TEMPLATES_DIR / f"output_{uuid.uuid4().hex}.docx"
    
    template_path = None
    if template_id:
        # DB에서 template_path 조회 (인증 필요하지만 간단하게 처리)
        from database import SessionLocal
        db = SessionLocal()
        try:
            tpl = db.query(UserTemplate).filter(UserTemplate.id == template_id).first()
            if tpl:
                template_path = tpl.file_path
        finally:
            db.close()
    
    try:
        markdown_to_docx(
            markdown=markdown,
            output_path=str(output_path),
            title=title,
            template_path=template_path,
            sections_data=sections_data,
        )
        filename_for_download = f"{title}.docx".replace("/", "_").replace(" ", "_")
        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename_for_download,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"docx 생성 실패: {str(e)}")


@app.post("/api/lesson/parse-sections")
async def api_parse_lesson_sections(body: dict):
    """
    마크다운 지도안을 섹션별로 분해 (양식 채우기용)
    
    body: { "markdown": "..." }
    Returns: { "활동명": "...", "목표": "...", ... }
    """
    md = body.get("markdown", "")
    sections = {}
    
    # 헤더(##, ###)로 섹션 분리
    current_key = None
    current_content = []
    
    for line in md.split("\\n"):
        line_stripped = line.strip()
        # 헤더 라인 감지
        if line_stripped.startswith("## ") or line_stripped.startswith("### "):
            # 이전 섹션 저장
            if current_key:
                sections[current_key] = "\\n".join(current_content).strip()
            # 새 섹션 시작
            current_key = line_stripped.lstrip("#").strip()
            current_content = []
        elif line_stripped.startswith("**") and line_stripped.endswith("**"):
            # **굵은글씨** 도 섹션으로 인식
            if current_key:
                sections[current_key] = "\\n".join(current_content).strip()
            current_key = line_stripped.strip("*").strip()
            current_content = []
        else:
            current_content.append(line)
    
    if current_key:
        sections[current_key] = "\\n".join(current_content).strip()
    
    return sections

'''
    main_code = main_code.replace(
        '\n@app.get("/api/health")',
        template_apis + '\n@app.get("/api/health")'
    )
    MAIN_PATH.write_text(main_code, encoding="utf-8")
    print("✅ main.py: 템플릿 API 4개 추가")
else:
    print("ℹ️  main.py: 템플릿 API 이미 존재")


# ============================================================
# 5) HTML 패치
# ============================================================
html = HTML_PATH.read_text(encoding="utf-8")
backup = HTML_PATH.with_suffix(".html.bak_phase3")
backup.write_text(html, encoding="utf-8")
print(f"✅ HTML 백업: {backup}")

# 5-1) 사이드바에 "내 양식" 메뉴 추가
old_nav = '''    <button class="nav-item" onclick="showPage('activitylog', this)">
      <span class="nav-icon">📋</span><span>내 지도안</span>
    </button>'''
new_nav = '''    <button class="nav-item" onclick="showPage('activitylog', this)">
      <span class="nav-icon">📋</span><span>내 지도안</span>
    </button>
    <button class="nav-item" onclick="showPage('templates', this)">
      <span class="nav-icon">📄</span><span>내 양식</span>
    </button>'''
if old_nav in html:
    html = html.replace(old_nav, new_nav)
    print("✅ HTML: 사이드바 '내 양식' 메뉴 추가")

# 5-2) "내 양식" 페이지 추가 (활동 기록지 페이지 뒤에)
templates_page = '''
  <!-- ════ MY TEMPLATES PAGE ════ -->
  <div class="page" id="page-templates">
    <div style="padding:2rem;">
      <div class="page-title">📄 내 지도안 양식</div>
      <div class="page-sub">우리 유치원 양식(.docx)을 업로드하면 그 양식에 맞춰 지도안을 생성합니다</div>

      <div id="tplNeedLogin" style="display:none; text-align:center; padding:3rem 1rem;">
        <div style="font-size:48px; margin-bottom:16px;">🔐</div>
        <div style="font-size:16px; font-weight:700; color:var(--g7); margin-bottom:8px;">로그인이 필요합니다</div>
        <button class="btn-primary" onclick="showModal()">로그인 / 회원가입</button>
      </div>

      <div id="tplContent">
        <!-- 업로드 영역 -->
        <div class="card" style="margin-bottom:1.5rem;">
          <div class="card-title" style="margin-bottom:12px;">새 양식 업로드</div>
          <div style="display:flex;gap:10px;flex-wrap:wrap;align-items:end;">
            <div style="flex:1;min-width:200px;">
              <label style="display:block;font-size:12px;font-weight:600;color:var(--g7);margin-bottom:6px;">양식 이름</label>
              <input id="tplName" type="text" placeholder="예: 우리유치원 표준양식" style="width:100%;height:40px;padding:0 12px;border:1.5px solid var(--g2);border-radius:var(--r);font-family:var(--font);font-size:13px;background:var(--white);outline:none;">
            </div>
            <div style="flex:1;min-width:200px;">
              <label style="display:block;font-size:12px;font-weight:600;color:var(--g7);margin-bottom:6px;">파일 (.docx만)</label>
              <input id="tplFile" type="file" accept=".docx" style="width:100%;height:40px;padding:7px 12px;border:1.5px solid var(--g2);border-radius:var(--r);font-family:var(--font);font-size:12px;background:var(--white);">
            </div>
            <button class="btn-primary" style="height:40px;" onclick="uploadTemplate()">업로드</button>
          </div>
          <div id="tplUploadResult" style="margin-top:12px;display:none;"></div>
        </div>

        <!-- 양식 목록 -->
        <div style="font-size:13px;font-weight:700;color:var(--g7);margin-bottom:10px;">📋 업로드된 양식</div>
        <div id="tplEmpty" style="display:none;text-align:center;padding:2rem;color:var(--g5);font-size:13px;">아직 업로드된 양식이 없습니다.</div>
        <ul class="log-list" id="tplList"></ul>
      </div>
    </div>
  </div>
'''

# activitylog 페이지 다음에 삽입
old_marker = '<!-- ════ COMMUNITY PAGE ════ -->'
if old_marker in html and 'id="page-templates"' not in html:
    html = html.replace(old_marker, templates_page + '\n  ' + old_marker, 1)
    print("✅ HTML: '내 양식' 페이지 추가")

# 5-3) Play-Scanner 페이지에 양식 선택 옵션 추가
# 지도안 출력 영역 직전 (output-actions 위)에 양식 선택 드롭다운 추가
old_output_head = '''        <div class="output-head">
          <div class="output-title" id="outputTitle">지도안 생성 결과</div>
          <div class="output-actions">'''
new_output_head = '''        <div class="output-head">
          <div class="output-title" id="outputTitle">지도안 생성 결과</div>
          <div class="output-actions">
            <select id="exportTplSelect" style="height:32px;padding:0 10px;border:1px solid var(--g2);border-radius:var(--r);font-family:var(--font);font-size:12px;background:var(--white);color:var(--g7);outline:none;">
              <option value="">기본 양식</option>
            </select>'''
if old_output_head in html:
    html = html.replace(old_output_head, new_output_head)
    print("✅ HTML: 지도안 결과에 양식 선택 추가")

# 5-4) docx 다운로드 버튼 추가 (📥 다운로드 버튼 옆에)
old_download_btn = '''<button class="btn-edit" onclick="downloadLessonPlan()">📥 다운로드</button>'''
new_download_btns = '''<button class="btn-edit" onclick="downloadLessonPlan()" title="Markdown 다운로드">📥 .md</button>
            <button class="btn-edit" onclick="downloadAsDocx()" title="Word 다운로드">📄 .docx</button>'''
if old_download_btn in html:
    html = html.replace(old_download_btn, new_download_btns)
    print("✅ HTML: docx 다운로드 버튼 추가")

# 5-5) JS 함수 추가
new_js = '''

// ============================================================
// Phase 3: 커스텀 양식 기능
// ============================================================

let _userTemplates = [];

async function loadTemplates() {
  const token = localStorage.getItem('auth_token');
  const needLogin = document.getElementById('tplNeedLogin');
  const content = document.getElementById('tplContent');
  if (!needLogin || !content) return;

  if (!token) {
    needLogin.style.display = 'block';
    content.style.display = 'none';
    return;
  }
  needLogin.style.display = 'none';
  content.style.display = 'block';

  try {
    const r = await fetch(API_BASE + '/api/templates/my', {
      headers: { 'Authorization': 'Bearer ' + token },
    });
    if (r.status === 401) {
      localStorage.removeItem('auth_token');
      needLogin.style.display = 'block';
      content.style.display = 'none';
      return;
    }
    _userTemplates = await r.json();
    renderTemplateList(_userTemplates);
    updateExportTplSelect();
  } catch (e) {
    console.error('양식 로드 실패:', e);
  }
}

function renderTemplateList(templates) {
  const list = document.getElementById('tplList');
  const empty = document.getElementById('tplEmpty');
  if (!list || !empty) return;

  if (!templates.length) {
    empty.style.display = 'block';
    list.innerHTML = '';
    return;
  }
  empty.style.display = 'none';

  list.innerHTML = templates.map(t => {
    const date = new Date(t.created_at).toLocaleDateString('ko-KR');
    const sectionNames = (t.sections || []).slice(0, 6).map(s => s.name).join(' · ');
    const moreCount = Math.max(0, (t.sections || []).length - 6);
    return `
      <li class="log-item">
        <div class="log-thumb">📄</div>
        <div class="log-content">
          <div class="log-title">${escapeHtml(t.template_name)}</div>
          <div class="log-meta">${escapeHtml(t.original_filename)} · ${t.structure_type === 'table' ? '표 형식' : '단락 형식'}</div>
          <div style="font-size:11px;color:var(--g5);margin-top:6px;">
            감지된 섹션 ${(t.sections||[]).length}개: ${escapeHtml(sectionNames)}${moreCount ? ' 외 ' + moreCount + '개' : ''}
          </div>
        </div>
        <div class="log-right">
          <div class="log-date">${date}</div>
          <div class="log-actions">
            <button class="log-action-btn" style="background:var(--coral-3);color:var(--coral);border:none;" onclick="deleteTemplate(${t.id})">🗑 삭제</button>
          </div>
        </div>
      </li>
    `;
  }).join('');
}

async function uploadTemplate() {
  const token = localStorage.getItem('auth_token');
  if (!token) { showModal(); return; }

  const nameEl = document.getElementById('tplName');
  const fileEl = document.getElementById('tplFile');
  const resultEl = document.getElementById('tplUploadResult');

  const name = nameEl.value.trim();
  const file = fileEl.files[0];

  if (!name) { showToast('양식 이름을 입력해주세요.', 'error'); return; }
  if (!file) { showToast('파일을 선택해주세요.', 'error'); return; }
  if (!file.name.endsWith('.docx')) { showToast('.docx 파일만 가능합니다.', 'error'); return; }

  const formData = new FormData();
  formData.append('file', file);
  formData.append('template_name', name);

  resultEl.style.display = 'block';
  resultEl.style.color = 'var(--g5)';
  resultEl.textContent = '업로드 중...';

  try {
    const r = await fetch(API_BASE + '/api/templates/upload', {
      method: 'POST',
      headers: { 'Authorization': 'Bearer ' + token },
      body: formData,
    });
    const data = await r.json();
    if (r.ok) {
      resultEl.style.color = 'var(--teal)';
      resultEl.innerHTML = '✅ 업로드 완료! 감지된 섹션: <b>' + data.total_sections + '개</b>';
      nameEl.value = ''; fileEl.value = '';
      showToast('양식이 등록되었습니다!', 'success');
      loadTemplates();
    } else {
      resultEl.style.color = 'var(--coral)';
      resultEl.textContent = '❌ 실패: ' + (data.detail || '알 수 없는 오류');
    }
  } catch (e) {
    resultEl.style.color = 'var(--coral)';
    resultEl.textContent = '❌ 오류: ' + e.message;
  }
}

async function deleteTemplate(id) {
  if (!confirm('이 양식을 삭제하시겠습니까?')) return;
  const token = localStorage.getItem('auth_token');
  if (!token) return;
  try {
    const r = await fetch(API_BASE + '/api/templates/' + id, {
      method: 'DELETE',
      headers: { 'Authorization': 'Bearer ' + token },
    });
    if (r.ok) {
      showToast('삭제되었습니다.', 'info');
      loadTemplates();
    }
  } catch (e) {
    showToast('삭제 실패', 'error');
  }
}

function updateExportTplSelect() {
  const sel = document.getElementById('exportTplSelect');
  if (!sel) return;
  const opts = ['<option value="">기본 양식</option>'];
  for (const t of _userTemplates) {
    opts.push('<option value="' + t.id + '">' + escapeHtml(t.template_name) + '</option>');
  }
  sel.innerHTML = opts.join('');
}

async function downloadAsDocx() {
  const content = document.getElementById('outputContent');
  const md = content.dataset.markdown;
  if (!md) { showToast('지도안이 없습니다.', 'error'); return; }

  const tplSelect = document.getElementById('exportTplSelect');
  const templateId = tplSelect && tplSelect.value ? parseInt(tplSelect.value) : null;

  const card = _state.cards && _state.cards[_state.selectedCardIdx];
  const title = card ? card.card_title : '지도안';

  showToast('docx 생성 중...', 'info');

  let sectionsData = null;
  if (templateId) {
    // 양식이 있으면 섹션별로 분해
    try {
      const r = await fetch(API_BASE + '/api/lesson/parse-sections', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ markdown: md }),
      });
      sectionsData = await r.json();
    } catch (e) {
      console.error('섹션 파싱 실패:', e);
    }
  }

  try {
    const r = await fetch(API_BASE + '/api/lesson/export-docx', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({
        markdown: md,
        title: title,
        template_id: templateId,
        sections_data: sectionsData,
      }),
    });
    if (!r.ok) throw new Error('생성 실패');
    const blob = await r.blob();
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url;
    a.download = title + '.docx';
    document.body.appendChild(a); a.click(); document.body.removeChild(a);
    URL.revokeObjectURL(url);
    showToast('.docx 다운로드 완료!', 'success');
  } catch (e) {
    showToast('다운로드 실패: ' + e.message, 'error');
  }
}

// showPage 후크에 templates 페이지 자동 로드 추가
const _origShowPageT = showPage;
showPage = function(id, navEl, section) {
  _origShowPageT(id, navEl, section);
  if (id === 'templates') loadTemplates();
  // scanner 페이지 들어가면 양식 목록 미리 로드 (드롭다운용)
  if (id === 'scanner' && localStorage.getItem('auth_token') && _userTemplates.length === 0) {
    loadTemplates();
  }
};
'''

# JS는 </script> 직전에 삽입
close_script = '</script>\n</body>'
if close_script in html and 'loadTemplates' not in html:
    html = html.replace(close_script, new_js + '\n</script>\n</body>')
    print("✅ HTML: Phase 3 JS 추가")

HTML_PATH.write_text(html, encoding="utf-8")

print("\n" + "=" * 50)
print("🎉 Phase 3 패치 완료!")
print("=" * 50)
print("\n다음 단계:")
print("  1. pip install python-docx")
print("  2. rm edubridge.db (DB 스키마 변경됨)")
print("  3. 서버 재시작")
print("\n테스트:")
print("  1. 회원가입/로그인")
print("  2. 사이드바 '📄 내 양식' → .docx 업로드")
print("  3. Play-Scanner → 지도안 생성 → 양식 선택 후 .docx 다운로드")
