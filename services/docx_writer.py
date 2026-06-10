"""
services/docx_writer.py
========================
지도안 마크다운을 .docx로 변환
(사용자 양식이 있으면 그 양식에 채우고, 없으면 기본 양식 사용)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional, Dict
import re


def markdown_to_docx(
    markdown: str,
    output_path: str,
    title: str = "지도안",
    template_path: Optional[str] = None,
    sections_data: Optional[Dict[str, str]] = None,
):
    """
    Markdown을 .docx로 저장.

    template_path가 있으면 그 양식의 표/단락에 sections_data를 채워서 저장.
    없으면 새 문서를 만들어서 마크다운을 단순 변환.
    """
    if template_path and sections_data:
        _fill_template(template_path, output_path, sections_data)
    else:
        _create_simple_docx(markdown, output_path, title)


def _create_simple_docx(markdown: str, output_path: str, title: str):
    """기본 양식: 마크다운 → 단순 docx"""
    doc = Document()

    # 제목
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 마크다운 줄 단위 처리
    for line in markdown.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        # 헤더 처리
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(16)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(14)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            r.bold = True
            r.font.size = Pt(12)
        # 리스트
        elif line.lstrip().startswith(("- ", "* ", "• ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        # 번호 리스트
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            doc.add_paragraph(text, style="List Number")
        # 굵은 글씨 처리 (**text**)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    doc.save(output_path)


def _add_formatted_text(paragraph, text: str):
    """**bold** 마크다운을 docx 굵은 글씨로"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _fill_template(template_path: str, output_path: str, sections_data: Dict[str, str]):
    """
    사용자 양식 docx를 열어서 각 섹션에 값을 채워 넣기.

    핵심 전략:
    - 라벨 셀(예: "지도교사", "수업일시", "학습 주제")이 발견되면
    - 옆/아래의 데이터 셀을 **무조건 덮어쓰기** (예시 내용 무시)
    - sections_data에 매칭되는 값이 있을 때만 덮어쓰기
    """
    doc = Document(template_path)
    filled_cells = set()  # 이미 채운 셀 추적

    # 1) 표 기반 채우기
    for table in doc.tables:
        rows = list(table.rows)
        for row_idx, row in enumerate(rows):
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                cell_id = (id(table), row_idx, col_idx)
                if cell_id in filled_cells:
                    continue

                cell_text = cell.text.strip().rstrip(":：").strip()
                if not cell_text:
                    continue

                # 이 셀이 "라벨"인지 판단 (짧은 텍스트 + sections_data와 매칭)
                if not _looks_like_label(cell_text):
                    continue

                value = _find_matching_value(cell_text, sections_data)
                if not value:
                    continue

                # 전략 1: 같은 행의 다음 셀 (예시 내용이 있어도 무조건 덮어쓰기)
                if col_idx + 1 < len(cells):
                    next_cell = cells[col_idx + 1]
                    next_id = (id(table), row_idx, col_idx + 1)
                    next_text = next_cell.text.strip()
                    # 다음 셀이 또 다른 라벨이 아니면 (= 데이터 셀이면) 덮어쓰기
                    if next_id not in filled_cells and not _looks_like_label(next_text):
                        _replace_cell_text(next_cell, value)
                        filled_cells.add(next_id)
                        # 같은 값이 병합되어 옆 셀에도 있으면 그것도 채우기
                        merged_idx = col_idx + 2
                        while merged_idx < len(cells) and cells[merged_idx].text.strip() == next_text:
                            _replace_cell_text(cells[merged_idx], value)
                            filled_cells.add((id(table), row_idx, merged_idx))
                            merged_idx += 1
                        continue

                # 전략 2: 바로 아래 행의 같은 위치 셀
                if row_idx + 1 < len(rows):
                    below_cells = rows[row_idx + 1].cells
                    if col_idx < len(below_cells):
                        below_cell = below_cells[col_idx]
                        below_id = (id(table), row_idx + 1, col_idx)
                        below_text = below_cell.text.strip()
                        if below_id not in filled_cells and not _looks_like_label(below_text):
                            _replace_cell_text(below_cell, value)
                            filled_cells.add(below_id)

    # 2) 단락 기반 채우기
    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        text = para.text.strip().rstrip(":：").strip()
        if not text or not _looks_like_label(text):
            continue
        value = _find_matching_value(text, sections_data)
        if value and i + 1 < len(paragraphs):
            next_para = paragraphs[i + 1]
            # 단락은 비어있을 때만 (덮어쓰면 본문 망가짐)
            if not next_para.text.strip() or _is_filler(next_para.text):
                next_para.add_run(value)

    doc.save(output_path)


SECTION_LABEL_KEYWORDS = [
    "활동명", "주제", "제목", "단원", "차시",
    "대상", "연령", "인원", "장소", "일시", "날짜",
    "목표", "목적", "성취",
    "준비물", "교재", "교구", "재료",
    "도입", "전개", "정리", "마무리", "활동내용",
    "활동", "수업", "지도",
    "평가", "유의", "참고", "비고",
    "교사", "유아", "역할", "학생",
    "학교", "학년", "반", "학과",
    "단계", "내용", "형태",
]


def _looks_like_label(text: str) -> bool:
    """이 텍스트가 양식의 라벨(헤더)처럼 보이는지 판단"""
    text = text.strip()
    if not text or len(text) > 30:
        return False
    # 줄바꿈이 적고 짧으면 라벨일 가능성 높음
    if "\n" in text:
        # 줄바꿈이 있어도 각 줄이 짧으면 라벨 ("학습\n단계" 같은 경우)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        if all(len(l) <= 10 for l in lines):
            for kw in SECTION_LABEL_KEYWORDS:
                if kw in text:
                    return True
            return False
        return False
    # 콜론으로 끝나면 거의 확실히 라벨
    if text.endswith(":") or text.endswith("："):
        return True
    # 키워드 매칭
    for kw in SECTION_LABEL_KEYWORDS:
        if kw in text:
            return True
    return False


def _replace_cell_text(cell, value: str):
    """셀의 기존 텍스트를 모두 제거하고 새 값 삽입"""
    # 첫 단락만 남기고 나머지 단락 제거
    for para in cell.paragraphs[1:]:
        p_element = para._element
        p_element.getparent().remove(p_element)
    # 첫 단락의 모든 run 제거
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run.text = ""
    # 새 텍스트 삽입 (줄바꿈 처리)
    lines = value.split("\n")
    if lines:
        first_para.add_run(lines[0])
        for extra_line in lines[1:]:
            new_para = cell.add_paragraph()
            new_para.add_run(extra_line)


def _find_matching_value(section_name: str, sections_data: Dict[str, str]) -> Optional[str]:
    """섹션 이름과 매칭되는 값 찾기 (부분 매칭 허용)"""
    if not section_name:
        return None
    for key, value in sections_data.items():
        if key in section_name or section_name in key:
            return value
    return None


def _is_filler(text: str) -> bool:
    text = text.strip()
    return text in ["", "...", "___", "OOO", "○○○"] or text.startswith("예)")



def fill_template_by_positions(template_path: str, output_path: str, position_fills: dict):
    """
    셀 위치 기반으로 양식 채우기.

    Args:
        position_fills: {(table_idx, row, col): "content", ...}
    """
    doc = Document(template_path)

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                key = (table_idx, row_idx, col_idx)
                if key in position_fills:
                    content = position_fills[key]
                    if content:
                        _replace_cell_text(cell, str(content))

    doc.save(output_path)


def markdown_to_docx_with_analysis(
    markdown: str,
    output_path: str,
    title: str,
    template_path: str,
    template_analysis: dict,
    lesson_meta: dict = None,
    personal_info: dict = None,
):
    """
    심층 분석된 양식 정보로 docx 생성
    """
    from services.template_analyzer import fill_cells_with_lesson

    position_fills = fill_cells_with_lesson(
        template_analysis=template_analysis,
        lesson_markdown=markdown,
        lesson_meta=lesson_meta,
        personal_info=personal_info,
    )

    print(f"📊 채울 셀 위치: {len(position_fills)}개")

    fill_template_by_positions(template_path, output_path, position_fills)
