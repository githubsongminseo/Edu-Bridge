"""
services/template_parser.py
============================
업로드된 .docx 양식을 파싱해서 섹션/필드 구조를 추출
"""
from docx import Document
from typing import List, Dict
import re


def parse_docx_template(file_path: str) -> Dict:
    """
    .docx 파일을 분석하여 구조 정보를 반환

    Returns:
        {
            "sections": [{"name": "활동명", "type": "text", "order": 0}, ...],
            "has_tables": True,
            "raw_text_sample": "...",
            "structure_type": "table" | "paragraph"
        }
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": f"파일 파싱 실패: {e}"}

    sections = []
    raw_lines = []

    # 1) 표(table) 기반 양식 분석
    has_tables = len(doc.tables) > 0
    if has_tables:
        order = 0
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                # 첫 셀이 헤더(label)인 경우가 많음
                for i, cell_text in enumerate(cells):
                    if cell_text and len(cell_text) < 30 and not _is_filler_text(cell_text):
                        # 헤더 후보 (짧은 텍스트, 콜론 끝, 한국어 키워드)
                        if _looks_like_section_header(cell_text):
                            sections.append({
                                "name": _clean_section_name(cell_text),
                                "type": "table_cell",
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "col_idx": i,
                                "order": order,
                            })
                            order += 1

    # 2) 단락 기반 양식 분석
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        raw_lines.append(text)
        if _looks_like_section_header(text) and not has_tables:
            sections.append({
                "name": _clean_section_name(text),
                "type": "paragraph",
                "order": len(sections),
            })

    # 중복 제거 (이름 기준)
    seen = set()
    unique_sections = []
    for s in sections:
        if s["name"] not in seen:
            seen.add(s["name"])
            unique_sections.append(s)

    return {
        "sections": unique_sections,
        "has_tables": has_tables,
        "raw_text_sample": "\n".join(raw_lines[:10]),
        "structure_type": "table" if has_tables else "paragraph",
        "total_sections": len(unique_sections),
    }


SECTION_KEYWORDS = [
    "활동명", "주제", "제목", "단원", "차시",
    "대상", "연령", "인원", "장소", "일시", "날짜",
    "목표", "목적", "학습목표", "교육목표", "수업목표",
    "준비물", "교재", "교구", "재료",
    "도입", "전개", "정리", "마무리", "활동내용",
    "활동", "수업", "지도", "전개과정",
    "평가", "유의점", "참고", "비고",
    "교사", "유아", "역할",
]


def _looks_like_section_header(text: str) -> bool:
    """섹션 헤더 같은 텍스트인지 판단"""
    text = text.strip()
    if not text or len(text) > 30:
        return False
    # 콜론으로 끝나는 짧은 텍스트
    if text.endswith(":") or text.endswith("："):
        return True
    # 키워드 매칭
    for kw in SECTION_KEYWORDS:
        if kw in text:
            return True
    return False


def _is_filler_text(text: str) -> bool:
    """예시 텍스트나 채워야 할 빈칸인지"""
    fillers = ["예)", "예시", "(예", "...", "___", "OOO", "○○○"]
    return any(f in text for f in fillers)


def _clean_section_name(text: str) -> str:
    """섹션 이름 정리"""
    text = text.strip()
    text = re.sub(r"[:：\s]+$", "", text)
    text = re.sub(r"^[\d\.\)\s]+", "", text)
    return text.strip()


def parse_hwp_template(file_path: str) -> dict:
    """
    .hwp 파일에서 텍스트를 추출하고 섹션을 감지.
    
    한계: .hwp는 표 구조를 정확히 파싱하기 어려워서 텍스트 기반으로만 분석.
    """
    text_lines = []
    
    # 방법 1: pyhwp(hwp5) 사이 라이브러리 시도
    try:
        from hwp5.proc import rawunicode
        from hwp5.xmlmodel import Hwp5File
        from hwp5.binmodel import ParaText
        
        with Hwp5File(file_path) as hwp:
            for paragraph in hwp.bodytext.section(0):
                if hasattr(paragraph, 'text'):
                    text = paragraph.text
                    if text and text.strip():
                        text_lines.append(text.strip())
    except Exception:
        # 방법 2: olefile로 raw 텍스트 추출
        try:
            import olefile
            import zlib
            
            ole = olefile.OleFileIO(file_path)
            
            # BodyText 스트림 읽기
            dirs = ole.listdir()
            for path in dirs:
                if "BodyText" in path[0] if path else False:
                    try:
                        stream = ole.openstream(path)
                        data = stream.read()
                        # HWP는 zlib 압축된 경우가 많음
                        try:
                            data = zlib.decompress(data, -15)
                        except:
                            pass
                        # 유니코드 텍스트 추출 시도
                        try:
                            text = data.decode("utf-16-le", errors="ignore")
                            # 제어 문자 제거
                            text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
                            for line in text.split("\n"):
                                line = line.strip()
                                if line and len(line) > 1:
                                    text_lines.append(line)
                        except Exception:
                            pass
                    except Exception:
                        continue
            ole.close()
        except Exception as e:
            return {"error": f".hwp 파일 파싱 실패: {e}. .docx로 변환 후 업로드해주세요."}
    
    if not text_lines:
        return {"error": "텍스트를 추출할 수 없습니다. .docx로 변환 후 업로드해주세요."}
    
    # 섹션 감지
    sections = []
    seen = set()
    order = 0
    for line in text_lines:
        if _looks_like_section_header(line) and not _is_filler_text(line):
            name = _clean_section_name(line)
            if name and name not in seen:
                seen.add(name)
                sections.append({
                    "name": name,
                    "type": "hwp_text",
                    "order": order,
                })
                order += 1
    
    return {
        "sections": sections,
        "has_tables": False,  # hwp는 표 구조 보존 불가
        "raw_text_sample": "\n".join(text_lines[:10]),
        "structure_type": "hwp_text",
        "total_sections": len(sections),
        "warning": "⚠️ .hwp는 텍스트만 추출됩니다. 출력은 .docx로만 가능합니다.",
    }
